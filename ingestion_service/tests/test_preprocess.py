from ingestion_service.preprocess import extract_text, _sha256
import pytest 
from pathlib import Path
import ingestion_service.preprocess as preprocess

def test_extract_txt(tmp_path):
    sample_txt = tmp_path / "example.txt"
    sample_txt.write_text("Hello\nWorld\n   Foo ")
    result = preprocess.extract_text(sample_txt)
    assert "text" in result
    assert "metadata" in result
    assert result["text"] == "Hello World Foo"
    assert result["metadata"]["source_file"] == "example.txt"
    assert result["metadata"]["char_count"] == len("Hello World Foo")
    assert isinstance(result["metadata"]["sha256"], str)

def test_extract_txt_not_found(tmp_path):
    sample = tmp_path / "missing.txt"
    with pytest.raises(FileNotFoundError):
        preprocess.extract_text(sample)

def test_extract_text_unsupported_type(tmp_path):
    weird_file = tmp_path / "foo.xlsx"
    weird_file.write_text("data")
    with pytest.raises(ValueError):
        preprocess.extract_text(weird_file)

def test_sha256(tmp_path):
    f = tmp_path / "file.txt"
    f.write_text("hello world")
    h1 = preprocess._sha256(f)
    h2 = preprocess._sha256(f)
    assert h1 == h2
    # Confirm it's 64 hex chars
    assert len(h1) == 64
    assert all(c in '0123456789abcdef' for c in h1)

def test_normalize():
    assert preprocess._normalize("A   B\n\nC") == "A B C"
    assert preprocess._normalize("   Many     spaces   here  ") == "Many spaces here"

def test_extract_docx(tmp_path):
    docx_path = tmp_path / "sample.docx"
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx is not installed")
    doc = Document()
    doc.add_paragraph("Hello Docx")
    doc.add_paragraph("Another Paragraph")
    doc.save(docx_path)
    result = preprocess.extract_text(docx_path)
    assert "Hello Docx" in result["text"]
    assert "Another Paragraph" in result["text"]
    assert result["metadata"]["source_file"] == "sample.docx"

def test_extract_pdf_native_and_ocr(monkeypatch, tmp_path):
    pdf_path = tmp_path / "sample.pdf"
    # simulate a PDF file (content does not matter for this test)
    pdf_path.write_bytes(b"%PDF-1.4\n%EOF\n")
    # patch _pdf_text to return some text
    monkeypatch.setattr(preprocess, "_pdf_text", lambda _: "X" * 100)
    monkeypatch.setattr(preprocess, "_pdf_ocr", lambda _: "OCRTEXT")
    result = preprocess.extract_text(pdf_path)
    assert result["text"] == "X" * 100
    # now patch native extraction to be poor (length < 50), triggers OCR
    monkeypatch.setattr(preprocess, "_pdf_text", lambda _: "")
    monkeypatch.setattr(preprocess, "_pdf_ocr", lambda _: "OCR Fallback Text")
    result2 = preprocess.extract_text(pdf_path)
    assert result2["text"] == "OCR Fallback Text"

def test_pdf_ocr_importerror(monkeypatch, tmp_path):
    pdf_path = tmp_path / "should_ocr.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n%EOF\n")
    # Patch out Imports to simulate missing OCR deps
    monkeypatch.setattr(preprocess, "convert_from_path", None, raising=False)
    import importlib
    orig_import = __import__

    def fake_import(name, *a, **k):
        if name in ("pytesseract", "pdf2image"):
            raise ImportError
        return orig_import(name, *a, **k)
    monkeypatch.setattr("builtins.__import__", fake_import)
    try:
        res = preprocess._pdf_ocr(pdf_path)
    except ImportError:
        res = "" 
    assert res == ""

